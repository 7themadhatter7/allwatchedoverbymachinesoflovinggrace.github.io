#!/usr/bin/env python3
"""
doc_to_triples.py
=================
Ghost in the Machine Labs

Standardized translator: architectural document → RM substrate format.

Pipeline:
  DOC → DECOMPOSE → CONCEPT GRAPH → VOCAB EXTENSION → RM TRIPLES → QUERY BATTERY

Produces three output files:
  {stem}_triples.json      — (subject, relation, object) concept graph
  {stem}_vocab.json        — new vocabulary entries in Edinburgh association format
  {stem}_queries.json      — structured query battery for RM feedback loop

Usage:
  python3 doc_to_triples.py document.docx [--out-dir ./output] [--verbose]

The triple format is directly compatible with RM's session memory consolidation
pipeline (05_session_memory.py). Load triples into RM via:
  extend_rm_vocab.py  →  loads vocab pairs into association memory
  load_concept_graph.py  →  loads triples as substrate training pairs
"""

import json
import re
import sys
import argparse
from pathlib import Path
from collections import defaultdict, Counter
from typing import Optional

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx required. Run: pip install python-docx --break-system-packages")
    sys.exit(1)


# ─── RELATION VOCABULARY ──────────────────────────────────────────────────────
# Fixed relation set. Every triple uses one of these.
# Extending this set requires a version bump.

RELATIONS = {
    # Structural
    "is_type":          "subject is an instance of object",
    "has_component":    "subject contains object as a sub-component",
    "depends_on":       "subject requires object to function",
    "produces":         "subject outputs object",
    "consumes":         "subject takes object as input",
    "transforms_to":    "subject state becomes object state",
    "implements":       "subject is a concrete realization of object",
    "extends":          "subject adds capability to object",
    # Behavioral
    "executes":         "subject runs object as a process",
    "validates":        "subject checks correctness of object",
    "stores":           "subject persists object",
    "loads":            "subject reads object into memory",
    "calls":            "subject invokes object",
    "returns":          "subject outputs object as result",
    # Descriptive
    "has_property":     "subject has attribute object",
    "has_constraint":   "subject is bounded by object",
    "precedes":         "subject must complete before object begins",
    "follows":          "subject begins after object completes",
    "maps_to":          "subject corresponds to object in another domain",
    "defined_by":       "subject is specified by object",
    # Governance
    "governed_by":      "subject decisions belong to object",
    "owned_by":         "subject is maintained by object",
}

# ─── CONCEPT NORMALIZER ───────────────────────────────────────────────────────

def normalize(text: str) -> str:
    """Normalize concept to snake_case token for substrate indexing."""
    text = text.strip().lower()
    # Remove punctuation except internal hyphens/underscores
    text = re.sub(r"[^\w\s\-]", "", text)
    # Collapse whitespace and hyphens to underscore
    text = re.sub(r"[\s\-]+", "_", text)
    # Remove leading/trailing underscores
    text = text.strip("_")
    return text


# ─── KNOWN CONCEPT MAP ────────────────────────────────────────────────────────
# Maps surface forms found in documents to canonical substrate tokens.
# Extend this as new documents introduce new surface forms.

CONCEPT_MAP = {
    # Architecture entities
    "rm": "resonant_mother",
    "the resonant mother": "resonant_mother",
    "mother": "resonant_mother",
    "e8 engine": "e8_arc_engine",
    "e8 arc engine": "e8_arc_engine",
    "e8 substrate": "e8_geometric_substrate",
    "geometric substrate": "e8_geometric_substrate",
    "e8 lattice": "e8_geometric_substrate",
    "sparky": "sparky_host",
    "arcy": "arcy_host",
    "arc-agi": "arc_agi_benchmark",
    "arc prize": "arc_prize",
    # Pipeline phases
    "phase 0": "pipeline_phase_0",
    "phase 1": "pipeline_phase_1",
    "phase 2": "pipeline_phase_2",
    "phase 3": "pipeline_phase_3",
    "phase 4": "pipeline_phase_4",
    "phase 5": "pipeline_phase_5",
    "phase 6": "pipeline_phase_6",
    "ingestion": "pipeline_phase_0",
    "domain triangulation": "pipeline_phase_1",
    "ambiguity resolution": "pipeline_phase_2",
    "archetypal matching": "pipeline_phase_3",
    "delta resolution": "pipeline_phase_4",
    "chunking": "pipeline_phase_5",
    "composition validation": "pipeline_phase_6",
    # Data structures
    "intent vector": "intent_vector",
    "state vector": "state_vector",
    "field matrix": "field_matrix",
    "f matrix": "field_matrix",
    "delta vector": "delta_vector",
    "concept graph": "concept_graph",
    "triple": "concept_triple",
    "triples": "concept_triple",
    # Functions
    "loose_to_structured": "fn_loose_to_structured",
    "triangulate_domains": "fn_triangulate_domains",
    "resolve_ambiguities": "fn_resolve_ambiguities",
    "match_archetypes": "fn_match_archetypes",
    "resolve_delta_to_components": "fn_resolve_delta_to_components",
    "chunk_for_fields": "fn_chunk_for_fields",
    "validate_composition": "fn_validate_composition",
    # Translation pipeline
    "doc_to_triples": "fn_doc_to_triples",
    "extend_rm_vocab": "fn_extend_rm_vocab",
    "load_concept_graph": "fn_load_concept_graph",
    "generate_queries": "fn_generate_queries",
    # Infrastructure
    "ollama": "ollama_backend",
    "tailscale": "tailscale_network",
    "ngrok": "ngrok_tunnel",
    "systemd": "systemd_service_manager",
    "association memory": "edinburgh_association_memory",
    "edinburgh associative thesaurus": "edinburgh_association_memory",
    # Concepts
    "llm": "large_language_model",
    "llm plugin": "llm_plugin",
    "pseudoinverse": "pseudoinverse_operation",
    "one-hot": "one_hot_encoding",
    "council": "governance_council",
    "divine mother edition": "divine_mother_edition",
}

def canonicalize(text: str) -> str:
    """Map surface form to canonical substrate token."""
    lower = text.strip().lower()
    if lower in CONCEPT_MAP:
        return CONCEPT_MAP[lower]
    return normalize(text)


# ─── DOCX PARSER ──────────────────────────────────────────────────────────────

class DocxParser:
    """
    Parse a .docx file into a structured intermediate representation.
    Preserves: heading hierarchy, body paragraphs, bullet lists, table rows, code blocks.
    """

    def __init__(self, path: str):
        self.doc = Document(path)
        self.sections = []      # [{title, level, paragraphs, tables, bullets}]
        self.current_section = None
        self._parse()

    def _heading_level(self, para) -> Optional[int]:
        if para.style is None: return None
        style = para.style.name
        for level in range(1, 5):
            if f"Heading {level}" in style:
                return level
        return None

    def _is_bullet(self, para) -> bool:
        if para.style is None: return False
        return para.style.name in ("List Bullet", "List Paragraph") or \
               para._element.find(qn("w:numPr")) is not None

    def _parse(self):
        current = {"title": "PREAMBLE", "level": 0, "paragraphs": [], "bullets": [], "tables": []}
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            level = self._heading_level(para)
            if level is not None:
                if current["paragraphs"] or current["bullets"] or current["tables"]:
                    self.sections.append(current)
                current = {"title": text, "level": level, "paragraphs": [], "bullets": [], "tables": []}
            elif self._is_bullet(para):
                current["bullets"].append(text)
            else:
                current["paragraphs"].append(text)
        
        if current["paragraphs"] or current["bullets"] or current["tables"]:
            self.sections.append(current)
        
        # Extract tables
        for table in self.doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(cells)
            if rows and self.sections:
                self.sections[-1]["tables"].append(rows)


# ─── TRIPLE EXTRACTORS ────────────────────────────────────────────────────────

class TripleExtractor:
    """
    Extract (subject, relation, object) triples from parsed document sections.
    
    Strategy per content type:
      Headings      → section anchor + hierarchy triples
      Body text     → pattern-matched relation extraction
      Bullets       → key:value parsing + is_type/has_property
      Tables        → header-keyed relation triples (row per mapping)
      Code blocks   → function signature extraction
    """

    # Regex patterns for relation extraction from prose
    PATTERNS = [
        # "X produces Y" / "X outputs Y"
        (r"(\w[\w\s]+?)\s+(?:produces?|outputs?|generates?|emits?)\s+([\w\s]+?)(?:\.|,|$)",
         "produces"),
        # "X consumes Y" / "X takes Y as input"
        (r"(\w[\w\s]+?)\s+(?:consumes?|takes?|ingests?|accepts?|receives?)\s+([\w\s]+?)(?:\s+as\s+input)?(?:\.|,|$)",
         "consumes"),
        # "X depends on Y" / "X requires Y"
        (r"(\w[\w\s]+?)\s+(?:depends?\s+on|requires?|needs?)\s+([\w\s]+?)(?:\.|,|$)",
         "depends_on"),
        # "X transforms to Y" / "X becomes Y"
        (r"(\w[\w\s]+?)\s+(?:transforms?\s+to|becomes?|maps?\s+to)\s+([\w\s]+?)(?:\.|,|$)",
         "transforms_to"),
        # "X validates Y" / "X checks Y"
        (r"(\w[\w\s]+?)\s+(?:validates?|verifies?|checks?)\s+([\w\s]+?)(?:\.|,|$)",
         "validates"),
        # "X is a/an Y" / "X is the Y"
        (r"(\w[\w\s]+?)\s+is\s+(?:a|an|the)\s+([\w\s]+?)(?:\.|,|$)",
         "is_type"),
        # "X stores Y" / "X persists Y"
        (r"(\w[\w\s]+?)\s+(?:stores?|persists?|saves?|holds?)\s+([\w\s]+?)(?:\.|,|$)",
         "stores"),
        # "X loads Y" / "X reads Y"
        (r"(\w[\w\s]+?)\s+(?:loads?|reads?)\s+([\w\s]+?)(?:\.|,|$)",
         "loads"),
        # "X calls Y" / "X invokes Y"
        (r"(\w[\w\s]+?)\s+(?:calls?|invokes?)\s+([\w\s]+?)(?:\.|,|$)",
         "calls"),
        # "X governed by Y"
        (r"(\w[\w\s]+?)\s+(?:governed\s+by|owned\s+by)\s+([\w\s]+?)(?:\.|,|$)",
         "governed_by"),
        # "before X" / "precedes Y" — ordering
        (r"(\w[\w\s]+?)\s+(?:precedes?|before|must\s+complete\s+before)\s+([\w\s]+?)(?:\.|,|$)",
         "precedes"),
    ]

    # Bullet patterns: "Label: description" or "Label — description"
    BULLET_SPLIT = re.compile(r"^(.+?)[\:\—\-]{1,2}\s+(.+)$")

    def __init__(self, verbose=False):
        self.verbose = verbose
        self.triples = []
        self.seen = set()
        self._section_stack = []  # Track heading hierarchy

    def _add(self, subject: str, relation: str, obj: str, source: str = ""):
        """Add triple with deduplication."""
        s = canonicalize(subject)
        o = canonicalize(obj)
        if not s or not o or s == o or len(s) < 2 or len(o) < 2:
            return
        # Filter noise tokens
        noise = {"the", "a", "an", "this", "that", "these", "those", "it", "its",
                 "they", "their", "which", "who", "where", "when", "how", "what",
                 "also", "only", "just", "all", "each", "every", "any", "no",
                 "not", "never", "always", "very", "more", "most", "less",
                 "used", "using", "used_for", "used_by", "example"}
        if s in noise or o in noise:
            return
        key = (s, relation, o)
        if key not in self.seen:
            self.seen.add(key)
            triple = {"subject": s, "relation": relation, "object": o}
            if source:
                triple["source"] = source
            self.triples.append(triple)
            if self.verbose:
                print(f"  TRIPLE: {s} --[{relation}]--> {o}")

    def extract_from_section(self, section: dict):
        title = section["title"]
        level = section["level"]
        title_token = canonicalize(title)

        # Update section stack — only pop if level > 0 to avoid infinite loop
        if level > 0:
            while len(self._section_stack) >= level and self._section_stack:
                self._section_stack.pop()
        self._section_stack.append(title_token)

        # Parent-child section hierarchy
        if len(self._section_stack) >= 2:
            parent = self._section_stack[-2]
            self._add(parent, "has_component", title_token, f"heading:{title}")

        # Extract from body paragraphs
        for para in section["paragraphs"]:
            self._extract_from_text(para, source=f"body:{title[:40]}")

        # Extract from bullets
        for bullet in section["bullets"]:
            self._extract_from_bullet(bullet, context=title_token, source=f"bullet:{title[:40]}")

        # Extract from tables
        for table in section["tables"]:
            self._extract_from_table(table, context=title_token)

    def _extract_from_text(self, text: str, source: str):
        """Pattern-match relations from prose."""
        for pattern, relation in self.PATTERNS:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                subj = match.group(1).strip()
                obj  = match.group(2).strip()
                # Filter overly long matches (noise)
                if len(subj.split()) <= 6 and len(obj.split()) <= 6:
                    self._add(subj, relation, obj, source)

    def _extract_from_bullet(self, text: str, context: str, source: str):
        """Extract from 'Label: description' bullet format."""
        m = self.BULLET_SPLIT.match(text)
        if m:
            label = m.group(1).strip()
            desc  = m.group(2).strip()
            label_token = canonicalize(label)
            # Label is a component of the context section
            self._add(context, "has_component", label_token, source)
            # Extract relations from the description
            self._extract_from_text(desc, source)
            # If description starts with a known concept, add is_type
            first_words = " ".join(desc.split()[:4])
            self._add(label_token, "defined_by", canonicalize(first_words), source)
        else:
            # Plain bullet — add as component of context
            words = text.split()
            if 2 <= len(words) <= 10:
                self._add(context, "has_component", canonicalize(text), source)
            self._extract_from_text(text, source)

    def _extract_from_table(self, table: list, context: str):
        """Extract from table rows. First row treated as headers if all caps or short."""
        if not table:
            return
        
        header_row = table[0]
        is_header = all(len(cell.split()) <= 4 for cell in header_row if cell)
        
        if is_header and len(table) > 1:
            headers = [canonicalize(h) for h in header_row]
            for row in table[1:]:
                if len(row) >= 2:
                    key = canonicalize(row[0])
                    for i, cell in enumerate(row[1:], 1):
                        if i < len(headers) and cell.strip():
                            relation = self._header_to_relation(headers[i])
                            self._add(key, relation, canonicalize(cell[:60]), f"table:{context[:30]}")
                            # Also extract prose from cell
                            if len(cell.split()) >= 3:
                                self._extract_from_text(cell, source=f"table:{context[:30]}")
        else:
            # No headers — treat as key-value pairs
            for row in table:
                if len(row) >= 2 and row[0] and row[1]:
                    self._add(canonicalize(row[0]), "defined_by", canonicalize(row[1][:60]),
                               f"table:{context[:30]}")

    def _header_to_relation(self, header: str) -> str:
        """Map table column header to a relation type."""
        mapping = {
            "specification": "defined_by",
            "description": "defined_by",
            "function": "executes",
            "status": "has_property",
            "requirement": "has_constraint",
            "output": "produces",
            "input": "consumes",
            "method": "implements",
            "domain": "maps_to",
            "content": "defined_by",
            "rule": "has_constraint",
            "evidence": "validated_by",
            "work_item": "has_component",
            "capability": "executes",
        }
        for key, rel in mapping.items():
            if key in header:
                return rel
        return "defined_by"

    # ─── CODE BLOCK EXTRACTION ───────────────────────────────────────────────

    def extract_function_signatures(self, text: str, source: str):
        """Extract function → parameter/return relationships from signature strings."""
        # Match: function_name(param: type, ...) -> return_type
        fn_pattern = re.compile(r"(\w+)\(([^)]*)\)\s*(?:->|→)\s*([\w\[\], ]+)")
        for match in fn_pattern.finditer(text):
            fn_name = canonicalize(match.group(1))
            params  = match.group(2)
            ret     = canonicalize(match.group(3).strip())
            self._add(fn_name, "returns", ret, source)
            for param in params.split(","):
                parts = param.strip().split(":")
                if parts:
                    pname = canonicalize(parts[0].strip())
                    self._add(fn_name, "consumes", pname, source)


# ─── VOCAB EXTRACTOR ──────────────────────────────────────────────────────────

class VocabExtractor:
    """
    Build Edinburgh-format association pairs from the concept graph.
    
    Edinburgh format:
      { "stimulus": "word", "response": "associated_word", "weight": float }
    
    Generates associations by:
      1. Co-occurrence in same section (weak association, weight 0.3)
      2. Direct triple relationship (strong association, weight 0.8)
      3. Shared relation type (medium association, weight 0.5)
    """

    def __init__(self, triples: list, sections: list):
        self.triples = triples
        self.sections = sections
        self.pairs = []
        self.seen = set()

    def _add_pair(self, stimulus: str, response: str, weight: float, source: str):
        key = (stimulus, response)
        if key not in self.seen and stimulus != response:
            self.seen.add(key)
            self.pairs.append({
                "stimulus": stimulus,
                "response": response,
                "weight": round(weight, 3),
                "source": source
            })

    def extract(self):
        # 1. Direct triple associations (highest weight)
        for t in self.triples:
            s, r, o = t["subject"], t["relation"], t["object"]
            self._add_pair(s, o, 0.8, f"triple:{r}")
            self._add_pair(o, s, 0.6, f"triple_rev:{r}")
            # Relation token itself as association
            self._add_pair(s, r, 0.4, "relation_name")
            self._add_pair(r, o, 0.4, "relation_name")

        # 2. Shared-subject co-occurrence
        by_subject = defaultdict(list)
        for t in self.triples:
            by_subject[t["subject"]].append(t["object"])
        for subj, objects in by_subject.items():
            for i, o1 in enumerate(objects):
                for o2 in objects[i+1:i+4]:  # limit to 3 co-occurrences
                    self._add_pair(o1, o2, 0.3, "co-object")

        # 3. Section co-occurrence (all concepts mentioned in same section)
        for section in self.sections:
            concepts = []
            for para in section["paragraphs"] + section["bullets"]:
                for surface, canonical in CONCEPT_MAP.items():
                    if surface in para.lower():
                        concepts.append(canonical)
            concepts = list(set(concepts))
            for i, c1 in enumerate(concepts):
                for c2 in concepts[i+1:i+5]:
                    self._add_pair(c1, c2, 0.25, f"section:{normalize(section['title'])[:30]}")

        return self.pairs


# ─── QUERY GENERATOR ──────────────────────────────────────────────────────────

class QueryGenerator:
    """
    Generate structured query battery for RM feedback.
    
    Query types:
      completeness  — every producer has a consumer
      dependency    — no undeclared dependencies
      interface     — output of phase N matches input of phase N+1
      gap           — concepts referenced but not defined
      structural    — circular dependency detection
    """

    def __init__(self, triples: list):
        self.triples = triples
        self.by_subject = defaultdict(list)
        self.by_object  = defaultdict(list)
        self.by_relation = defaultdict(list)
        for t in triples:
            self.by_subject[t["subject"]].append(t)
            self.by_object[t["object"]].append(t)
            self.by_relation[t["relation"]].append(t)

    def generate(self) -> list:
        queries = []
        queries += self._completeness_checks()
        queries += self._interface_checks()
        queries += self._gap_checks()
        queries += self._dependency_checks()
        queries += self._archetype_queries()
        return queries

    def _completeness_checks(self) -> list:
        """Every concept that produces something should have a consumer."""
        queries = []
        producers = {t["subject"] for t in self.by_relation.get("produces", [])}
        consumers = {t["object"]  for t in self.by_relation.get("consumes", [])}
        unclaimed = producers - consumers
        for p in sorted(unclaimed)[:10]:  # cap at 10
            queries.append({
                "type": "completeness",
                "question": f"Is the output of '{p}' consumed by any downstream component?",
                "subject": p,
                "relation": "produces",
                "expected": "has_consumer",
                "impact": "high"
            })
        return queries

    def _interface_checks(self) -> list:
        """Check phase pipeline interface consistency."""
        queries = []
        phases = [f"pipeline_phase_{i}" for i in range(7)]
        for i in range(len(phases) - 1):
            p_current = phases[i]
            p_next    = phases[i + 1]
            current_outputs = {t["object"] for t in self.by_subject.get(p_current, [])
                               if t["relation"] == "produces"}
            next_inputs     = {t["object"] for t in self.by_subject.get(p_next, [])
                               if t["relation"] == "consumes"}
            if current_outputs and next_inputs:
                overlap = current_outputs & next_inputs
                queries.append({
                    "type": "interface",
                    "question": f"Does the output of phase {i} ({', '.join(list(current_outputs)[:3])}) "
                                f"match the input of phase {i+1} ({', '.join(list(next_inputs)[:3])})?",
                    "subject": p_current,
                    "object": p_next,
                    "shared_contracts": list(overlap),
                    "impact": "critical"
                })
        return queries

    def _gap_checks(self) -> list:
        """Find concepts referenced as objects but never defined as subjects."""
        queries = []
        all_subjects = set(self.by_subject.keys())
        all_objects  = set(self.by_object.keys())
        undefined    = all_objects - all_subjects
        # Filter out likely noise (very short tokens, relation names)
        undefined = {c for c in undefined if len(c) > 4 and "_" in c or len(c) > 8}
        for concept in sorted(undefined)[:10]:
            queries.append({
                "type": "gap",
                "question": f"Is '{concept}' fully specified, or does it require a definition section?",
                "subject": concept,
                "relation": "undefined",
                "impact": "medium"
            })
        return queries

    def _dependency_checks(self) -> list:
        """Detect potential circular dependencies."""
        queries = []
        deps = defaultdict(set)
        for t in self.by_relation.get("depends_on", []):
            deps[t["subject"]].add(t["object"])
        
        # Simple two-cycle detection
        for a, a_deps in deps.items():
            for b in a_deps:
                if a in deps.get(b, set()):
                    queries.append({
                        "type": "structural",
                        "question": f"Circular dependency detected: '{a}' depends on '{b}' and vice versa. Is this intentional (feedback loop) or an architectural error?",
                        "subject": a,
                        "object": b,
                        "relation": "circular_dependency",
                        "impact": "critical"
                    })
        return queries

    def _archetype_queries(self) -> list:
        """High-level architectural questions for RM's lattice navigation."""
        return [
            {
                "type": "archetype",
                "question": "What is the nearest known architectural archetype to the overall pipeline structure?",
                "subject": "software_engineering_pipeline",
                "relation": "maps_to",
                "object": "archetype_library",
                "impact": "high"
            },
            {
                "type": "archetype",
                "question": "Is the LLM plugin boundary correctly placed, or should any plugin-handled function be geometric?",
                "subject": "llm_plugin",
                "relation": "boundary_check",
                "object": "e8_geometric_substrate",
                "impact": "high"
            },
            {
                "type": "archetype",
                "question": "Which component in the pipeline has the highest geometric distance from its nearest archetype?",
                "subject": "pipeline_phase_all",
                "relation": "max_delta",
                "object": "archetype_library",
                "impact": "medium"
            },
        ]


# ─── STATISTICS REPORTER ──────────────────────────────────────────────────────

def report_stats(triples, vocab, queries, sections):
    print("\n" + "="*60)
    print("TRANSLATION REPORT")
    print("="*60)
    print(f"  Document sections parsed : {len(sections)}")
    print(f"  Concept triples produced : {len(triples)}")
    print(f"  Unique subjects          : {len(set(t['subject'] for t in triples))}")
    print(f"  Unique objects           : {len(set(t['object'] for t in triples))}")
    print(f"  Relation types used      : {len(set(t['relation'] for t in triples))}")
    print(f"  Vocabulary pairs         : {len(vocab)}")
    print(f"  Query battery size       : {len(queries)}")
    print()
    
    rel_counts = Counter(t["relation"] for t in triples)
    print("  Top relations:")
    for rel, count in rel_counts.most_common(8):
        print(f"    {rel:<30} {count}")
    print()
    
    query_types = Counter(q["type"] for q in queries)
    print("  Query types:")
    for qt, count in query_types.items():
        print(f"    {qt:<20} {count}")
    
    critical = [q for q in queries if q.get("impact") == "critical"]
    if critical:
        print(f"\n  CRITICAL queries ({len(critical)}):")
        for q in critical:
            print(f"    [{q['type']}] {q['question'][:80]}")
    print("="*60 + "\n")


# ─── MAIN ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Translate architectural doc to RM substrate format")
    parser.add_argument("document", help="Path to .docx file")
    parser.add_argument("--out-dir", default=".", help="Output directory")
    parser.add_argument("--verbose", action="store_true", help="Print each triple as extracted")
    args = parser.parse_args()

    doc_path = Path(args.document)
    if not doc_path.exists():
        print(f"ERROR: File not found: {doc_path}")
        sys.exit(1)

    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    stem = doc_path.stem

    print(f"\ndoc_to_triples.py — Ghost in the Machine Labs")
    print(f"Translating: {doc_path.name}")
    print(f"Output dir:  {out_dir}\n")

    # Parse document
    print("Step 1: Parsing document structure...")
    doc_parser = DocxParser(str(doc_path))
    print(f"  {len(doc_parser.sections)} sections found")

    # Extract triples
    print("Step 2: Extracting concept triples...")
    extractor = TripleExtractor(verbose=args.verbose)
    for section in doc_parser.sections:
        extractor.extract_from_section(section)
        # Also scan for function signatures in all text
        for para in section["paragraphs"]:
            extractor.extract_function_signatures(para, f"sig:{section['title'][:30]}")
    
    triples = extractor.triples
    print(f"  {len(triples)} triples extracted")

    # Build vocabulary
    print("Step 3: Building vocabulary extension...")
    vocab_extractor = VocabExtractor(triples, doc_parser.sections)
    vocab = vocab_extractor.extract()
    print(f"  {len(vocab)} association pairs generated")

    # Generate query battery
    print("Step 4: Generating RM query battery...")
    query_gen = QueryGenerator(triples)
    queries = query_gen.generate()
    print(f"  {len(queries)} queries generated")

    # Write outputs
    triples_path = out_dir / f"{stem}_triples.json"
    vocab_path   = out_dir / f"{stem}_vocab.json"
    queries_path = out_dir / f"{stem}_queries.json"

    with open(triples_path, "w") as f:
        json.dump({
            "source_document": str(doc_path.name),
            "concept_map_version": "1.0",
            "relation_vocabulary": list(RELATIONS.keys()),
            "triple_count": len(triples),
            "triples": triples
        }, f, indent=2)

    with open(vocab_path, "w") as f:
        json.dump({
            "source_document": str(doc_path.name),
            "format": "edinburgh_association",
            "pair_count": len(vocab),
            "pairs": vocab
        }, f, indent=2)

    with open(queries_path, "w") as f:
        json.dump({
            "source_document": str(doc_path.name),
            "query_count": len(queries),
            "queries": queries
        }, f, indent=2)

    report_stats(triples, vocab, queries, doc_parser.sections)

    print(f"Output files:")
    print(f"  Triples  → {triples_path}")
    print(f"  Vocab    → {vocab_path}")
    print(f"  Queries  → {queries_path}")
    print()
    print("Next steps:")
    print("  1. extend_rm_vocab.py    — load vocab pairs into RM association memory")
    print("  2. load_concept_graph.py — load triples as RM substrate training pairs")
    print("  3. run_rm_queries.py     — execute query battery against RM, collect feedback")


if __name__ == "__main__":
    main()
